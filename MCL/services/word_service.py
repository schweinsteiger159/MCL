from models.printable_row import PrintableRow
from models.printable_page import PrintablePage

from docx import Document
from docx.shared import Inches
from docx.shared import Pt

from models.printable_row import PrintableRow
from models.printable_page import PrintablePage

from io import BytesIO

from docx.shared import Inches
from docx.enum.section import WD_ORIENT

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Pt
from docx.shared import Cm

class WordService:

    # ==================================================
    # CONVERT CARTON -> PRINTABLE ROWS
    # ==================================================

    def build_printable_rows(self, carton):

        rows = []

        for style_group in carton.style_groups:

            first_color = True

            for color_group in style_group.color_groups:

                first_size = True

                size_count = len(
                    color_group.size_qty_list
                )

                for index, size_qty in enumerate(
                    color_group.size_qty_list
                ):

                    row = PrintableRow(

                        style=style_group.style,

                        color=color_group.color,

                        size=size_qty.size,

                        qty=size_qty.qty,

                        show_style=first_color and first_size,

                        show_color=first_size,

                        is_style_end=(
                            index == size_count - 1
                        )
                    )

                    rows.append(row)

                    first_size = False

                first_color = False

        return rows

    # ==================================================
    # PAGINATION
    # ==================================================

    def paginate_rows(
        self,
        carton,
        rows,
        rows_per_page=8
    ):

        pages = []

        total_rows = len(rows)

        total_pages = (
            total_rows + rows_per_page - 1
        ) // rows_per_page

        for page_index in range(total_pages):

            page = PrintablePage(
                carton=carton,
                page_no=page_index + 1,
                total_pages=total_pages
            )

            start = page_index * rows_per_page

            end = start + rows_per_page

            page_rows = rows[start:end]

            # ==================================
            # HANDLE CONTINUATION
            # ==================================

            if page_index > 0 and len(page_rows) > 0:

                
                page_rows[0].show_style = True
                page_rows[0].show_color = True


            page.rows = page_rows

            pages.append(page)

        return pages

    # ==================================================
    # EXPORT WORD
    # ==================================================

    def export_word(
        self,
        cartons
    ):

        document = Document()
        # ==========================================
        # PAGE SETUP
        # ==========================================

        section = document.sections[0]

        # A6 Landscape
        section.page_width = Inches(5.83)
        section.page_height = Inches(4.13)

        section.orientation = WD_ORIENT.LANDSCAPE

        # Margins
        section.top_margin = Inches(0.1)

        section.left_margin = Inches(0.1)

        section.right_margin = Inches(0)

        section.bottom_margin = Inches(0)

        for carton in cartons.values():

            rows = self.build_printable_rows(
                carton
            )

            pages = self.paginate_rows(
                carton,
                rows,
                rows_per_page=8
            )

            for page in pages:

                # ==========================================
                # HEADER
                # ==========================================

                p = document.add_paragraph()
                self.apply_paragraph_style(p)

                run = p.add_run(
                    f"PO: {carton.po}"
                )

                run.bold = True

                run.font.size = Pt(10)

                p = document.add_paragraph()
                self.apply_paragraph_style(p)

                run = p.add_run(
                    f"CTN: {carton.carton_no} of {carton.of}"
                )

                run.bold = True

                run.font.size = Pt(10)

                p = document.add_paragraph()
                self.apply_paragraph_style(p)

                run = p.add_run(
                    f"WEIGHT: {carton.weight} KG"
                )

                run.bold = True

                run.font.size = Pt(10)

                # ==========================================
                # BARCODE
                # ==========================================

                p = document.add_paragraph()

                self.apply_paragraph_style(p)


                run = p.add_run(
                    f"*{carton.case_id}*"
                )

                run.font.name = "Code 128"

                run.font.size = Pt(28)

                # CASE ID

                p = document.add_paragraph()
                self.apply_paragraph_style(p)

                run = p.add_run(
                    carton.case_id
                )

                run.bold = True

                run.font.size = Pt(11)


                

                # ==============================
                # TABLE
                # ==============================

                table = document.add_table(
                    rows=1,
                    cols=4
                )
                table.autofit = False
                table.columns[0].width = Inches(2.3)
                table.columns[1].width = Inches(1.3)
                table.columns[2].width = Inches(0.8)
                table.columns[3].width = Inches(0.8)

                #table.style = "Table Grid"

                header = table.rows[0].cells
                table.rows[0].height = Cm(0.5)

                header[0].text = "Style"
                header[1].text = "Color"
                header[2].text = "Size"
                header[3].text = "Qty"

                for cell in header:

                    for paragraph in cell.paragraphs:

                            paragraph.paragraph_format.space_after = Pt(3)

                            for run in paragraph.runs:

                                run.bold = True

                                run.font.size = Pt(11)

                for row in page.rows:

                    cells = table.add_row().cells
                    row_cells = table.rows[-1]

                    row_cells.height = Cm(0.45)
                    # ==========================================
                    # SET TEXT
                    # ==========================================

                    cells[0].text = (
                        row.style
                        if row.show_style
                        else ""
                    )

                    cells[1].text = (
                        row.color
                        if row.show_color
                        else ""
                    )

                    cells[2].text = str(row.size)

                    cells[3].text = str(row.qty)

                    # ==========================================
                    # FORMAT CELL
                    # ==========================================

                    for cell in cells:

                        for paragraph in cell.paragraphs:

                            paragraph.paragraph_format.space_after = Pt(3)

                            for run in paragraph.runs:

                                run.font.size = Pt(11)

                # ==============================================
                # PAGE FOOTER
                # ==============================================
                if page.total_pages > 1:

                    # push footer xuống thấp hơn
                    for _ in range(2):

                        p = document.add_paragraph()

                        self.apply_paragraph_style(p)

                    p = document.add_paragraph()

                    p.alignment = (
                        WD_PARAGRAPH_ALIGNMENT.CENTER
                    )

                    run = p.add_run(
                        f"{page.page_no} of "
                        f"{page.total_pages}"
                    )

                    run.font.size = Pt(10)

                # ==============================================
                # PAGE BREAK
                # ==============================================

                if not (
                    carton == list(cartons.values())[-1]
                    and
                    page.page_no == page.total_pages
                ):
                    document.add_page_break()

        # ==============================================
        # EXPORT BUFFER
        # ==============================================

        buffer = BytesIO()

        document.save(buffer)

        buffer.seek(0)

        return buffer

        # ==================================================
        # APPLY PARAGRAPH STYLE
        # ==================================================

    def apply_paragraph_style(
        self,
        paragraph
    ):

        paragraph.paragraph_format.space_after = Pt(3)
